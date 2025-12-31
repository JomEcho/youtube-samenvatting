"""
YouTube Samenvatting Tool
Haalt transcripties op van YouTube videos en maakt samenvattingen.
"""

import os
import re
import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Optional, Tuple
from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api._errors import TranscriptsDisabled, NoTranscriptFound
import requests
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Setup logging
LOG_FILE = Path.home() / ".youtube_samenvatting.log"
logging.basicConfig(
    filename=LOG_FILE,
    level=logging.ERROR,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# Load .env file if it exists
try:
    from dotenv import load_dotenv
    # Look for .env in the script directory
    env_path = Path(__file__).parent / ".env"
    if env_path.exists():
        load_dotenv(env_path)
except ImportError:
    pass

# Output directory
OUTPUT_DIR = Path.home() / "Documents" / "YouTube-Samenvattingen"

# Transcript limieten per provider (karakters) - dit zijn TOTALE context limieten
# De effectieve transcript limiet = totaal - prompt lengte - output buffer
TRANSCRIPT_LIMITS = {
    "ollama_gpt-oss_summary": 50000,   # gpt-oss:20b heeft groter context window
    "ollama_gpt-oss_chat": 30000,
    "ollama_gemma2_summary": 20000,    # gemma2:9b heeft 8K context (~30 min video)
    "ollama_gemma2_chat": 12000,
    "openai_summary": 100000,          # GPT-4o-mini heeft 128k context
    "openai_chat": 80000,
    "anthropic_summary": 150000,       # Claude heeft 200k context
    "anthropic_chat": 120000,
}

# Buffer voor prompt overhead en output
PROMPT_OVERHEAD = 3000  # ~prompt lengte + wat buffer


def get_effective_limit(base_limit: int, prompt_length: int = 0) -> int:
    """Bereken effectieve transcript limiet minus prompt overhead."""
    effective = base_limit - PROMPT_OVERHEAD - prompt_length
    return max(effective, 1000)  # Minimaal 1000 karakters

# Samenvatting prompt - "Granulaire Systeem-Analist"
SUMMARY_PROMPT = """BELANGRIJK: Schrijf de VOLLEDIGE samenvatting in het NEDERLANDS.

Rol: Treed op als een Senior Technical Lead en Systeemarchitect. Je doel is om een samenvatting te maken die maximale informatiedichtheid combineert met technische precisie.

Gouden Regel voor Specificiteit:
Vermijd vage generalisaties (zoals "men bespreekt AI-modellen" of "er is vooruitgang"). Gebruik in plaats daarvan de exacte eigennamen, versienummers, tools, bibliotheken en wetenschappelijke parameters die in de video worden genoemd. Als er wordt gesproken over "Claude Code in VS Code om een agent te bouwen", noteer dan exact die combinatie.

Hanteer deze structuur:

## Core Thesis (Het Fundament)
De essentie van de video in één technische stelling.

## Technische Componenten & Toolstack
Maak een lijst van alle specifieke tools, modellen, API's of wetenschappelijke ontdekkingen die zijn genoemd (bijv. AlphaFold 3, PyTorch, GPT-4o, GNoME-database). Beschrijf kort hun specifieke rol.

## Concrete Toepassingen (The 'How-To')
Beschrijf de exacte workflows of implementaties die zijn besproken.
- Slecht voorbeeld: "Ze bouwen apps met AI."
- Goed voorbeeld: "Gebruik van Claude Code CLI binnen een VS Code-omgeving voor autonoom refactoren van legacy Python-code."

## Mechanische Diepgang
Leg de onderliggende logica uit. Hoe werkt het proces precies? Wat zijn de beperkingen of 'bottlenecks' die zijn genoemd?

## Toekomstige Implicaties & 'Next Steps'
Wat zijn de directe gevolgen voor het vakgebied? Welke concrete voorspellingen worden er gedaan voor de komende 6-12 maanden?

---
Negatieve Constraints (Wat NIET te doen):
- Geen grappen, bantering of introductiepraatjes opnemen
- Geen vage werkwoorden zoals "bespreken", "onderzoeken" of "vinden" zonder direct object
- Geen metaforen tenzij ze essentieel zijn voor de technische uitleg
"""


def extract_video_id(url: str) -> Optional[str]:
    """Extract video ID from various YouTube URL formats."""
    patterns = [
        r'(?:youtube\.com\/watch\?v=|youtu\.be\/|youtube\.com\/embed\/)([a-zA-Z0-9_-]{11})',
        r'^([a-zA-Z0-9_-]{11})$'  # Direct video ID
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None


def get_video_title(video_id: str) -> str:
    """Get video title from YouTube (basic method without API key)."""
    try:
        url = f"https://www.youtube.com/watch?v={video_id}"
        response = requests.get(url, timeout=10)
        match = re.search(r'<title>(.+?) - YouTube</title>', response.text)
        if match:
            title = match.group(1)
            # Clean title for filename
            return re.sub(r'[<>:"/\\|?*]', '', title)[:100]
    except Exception:
        pass
    return f"video_{video_id}"


def get_transcript(video_id: str) -> Tuple[str, str]:
    """
    Get transcript from YouTube video.
    Returns (transcript_text, language)
    """
    api = YouTubeTranscriptApi()

    try:
        # Get list of available transcripts
        transcript_list = list(api.list(video_id))

        if not transcript_list:
            raise Exception("Geen transcripties beschikbaar voor deze video.")

        # Priority order for languages
        priority = ['nl', 'en']
        selected = None
        lang = "unknown"

        # First try to find preferred language
        for pref_lang in priority:
            for t in transcript_list:
                if t.language_code.startswith(pref_lang):
                    selected = t
                    lang = t.language_code
                    break
            if selected:
                break

        # If no preferred language found, take first available
        if not selected:
            selected = transcript_list[0]
            lang = selected.language_code

        # Fetch the transcript
        data = api.fetch(video_id, languages=[lang])
        full_text = "\n".join([entry.text for entry in data])
        return full_text, lang

    except TranscriptsDisabled:
        logging.warning(f"Transcripties uitgeschakeld voor video {video_id}")
        raise Exception("Transcripties zijn uitgeschakeld voor deze video.")
    except NoTranscriptFound:
        logging.warning(f"Geen transcriptie gevonden voor video {video_id}")
        raise Exception("Geen transcriptie gevonden voor deze video.")
    except Exception as e:
        logging.error(f"Fout bij ophalen transcriptie voor {video_id}", exc_info=True)
        raise Exception(f"Fout bij ophalen transcriptie: {type(e).__name__}: {str(e)}")


def summarize_with_ollama(text: str, model: str = "gpt-oss:20b") -> str:
    """Summarize text using local Ollama."""
    # Use model-specific limit, minus prompt overhead
    if "gemma" in model.lower():
        base_limit = TRANSCRIPT_LIMITS["ollama_gemma2_summary"]
    else:
        base_limit = TRANSCRIPT_LIMITS["ollama_gpt-oss_summary"]

    effective_limit = get_effective_limit(base_limit, len(SUMMARY_PROMPT))
    truncated_text = text[:effective_limit]

    prompt = f"""{SUMMARY_PROMPT}

---
TRANSCRIPTIE:
{truncated_text}
"""

    try:
        response = requests.post(
            "http://localhost:11434/api/generate",
            json={
                "model": model,
                "prompt": prompt,
                "stream": False,
                "options": {
                    "temperature": 0.3,
                    "num_predict": 2000
                }
            },
            timeout=300  # 5 minutes timeout for local model
        )
        response.raise_for_status()
        return response.json()["response"]
    except requests.exceptions.ConnectionError:
        logging.error("Kan geen verbinding maken met Ollama")
        raise Exception("Kan geen verbinding maken met Ollama. Is Ollama actief?")
    except requests.exceptions.Timeout:
        logging.error(f"Ollama timeout voor model {model}")
        raise Exception(f"Ollama timeout - het model {model} reageert niet binnen 5 minuten.")
    except Exception as e:
        logging.error(f"Ollama fout met model {model}", exc_info=True)
        raise Exception(f"Ollama fout: {type(e).__name__}: {str(e)}")


def summarize_with_openai(text: str, api_key: str) -> str:
    """Summarize text using OpenAI API."""
    from openai import OpenAI

    client = OpenAI(api_key=api_key)
    effective_limit = get_effective_limit(TRANSCRIPT_LIMITS["openai_summary"], len(SUMMARY_PROMPT))
    truncated_text = text[:effective_limit]

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": SUMMARY_PROMPT},
                {"role": "user", "content": f"TRANSCRIPTIE:\n{truncated_text}"}
            ],
            temperature=0.3,
            max_tokens=4000
        )
        return response.choices[0].message.content
    except Exception as e:
        logging.error("OpenAI API fout", exc_info=True)
        raise Exception(f"OpenAI fout: {type(e).__name__}: {str(e)}")


def summarize_with_anthropic(text: str, api_key: str) -> str:
    """Summarize text using Anthropic API."""
    import anthropic

    client = anthropic.Anthropic(api_key=api_key)
    effective_limit = get_effective_limit(TRANSCRIPT_LIMITS["anthropic_summary"], len(SUMMARY_PROMPT))
    truncated_text = text[:effective_limit]

    try:
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            messages=[
                {
                    "role": "user",
                    "content": f"{SUMMARY_PROMPT}\n\n---\nTRANSCRIPTIE:\n{truncated_text}"
                }
            ]
        )
        return response.content[0].text
    except Exception as e:
        logging.error("Anthropic API fout", exc_info=True)
        raise Exception(f"Anthropic fout: {type(e).__name__}: {str(e)}")


def summarize(text: str, provider: str, api_key: Optional[str] = None, model: str = None) -> str:
    """Summarize text using specified provider."""
    if provider == "ollama":
        return summarize_with_ollama(text, model or "gpt-oss:20b")
    elif provider == "openai":
        if not api_key:
            raise Exception("OpenAI API key is vereist.")
        return summarize_with_openai(text, api_key)
    elif provider == "anthropic":
        if not api_key:
            raise Exception("Anthropic API key is vereist.")
        return summarize_with_anthropic(text, api_key)
    else:
        raise Exception(f"Onbekende provider: {provider}")


# Chat system prompt - strikt gebaseerd op transcript
CHAT_SYSTEM_PROMPT = """Je bent een Nederlandstalige assistent die vragen beantwoordt over een YouTube video.
Je hebt ALLEEN toegang tot het transcript hieronder.

STRIKTE REGELS:
- Antwoord ALTIJD in het Nederlands
- Baseer je antwoord UITSLUITEND op het transcript
- Als het antwoord niet in het transcript staat, zeg: "Dit staat niet in de video."
- Citeer relevante passages uit het transcript waar mogelijk
- Verzin NOOIT informatie die niet in het transcript staat
- Als je onzeker bent, geef dat aan

TRANSCRIPT:
{transcript}"""


def chat_with_ollama(transcript: str, question: str, chat_history: list, model: str = "gpt-oss:20b") -> str:
    """Chat about transcript using local Ollama."""
    # Use model-specific limit, minus prompt overhead
    if "gemma" in model.lower():
        base_limit = TRANSCRIPT_LIMITS["ollama_gemma2_chat"]
    else:
        base_limit = TRANSCRIPT_LIMITS["ollama_gpt-oss_chat"]

    # Chat prompt is kleiner dan summary prompt
    chat_prompt_base_len = len(CHAT_SYSTEM_PROMPT) - len("{transcript}")
    effective_limit = get_effective_limit(base_limit, chat_prompt_base_len)
    system_prompt = CHAT_SYSTEM_PROMPT.format(transcript=transcript[:effective_limit])

    # Build conversation context
    messages_text = f"{system_prompt}\n\n"
    for msg in chat_history[-10:]:  # Last 10 messages for context
        role = "Gebruiker" if msg["role"] == "user" else "Assistent"
        messages_text += f"{role}: {msg['content']}\n\n"
    messages_text += f"Gebruiker: {question}\n\nAssistent:"

    try:
        response = requests.post(
            "http://localhost:11434/api/generate",
            json={
                "model": model,
                "prompt": messages_text,
                "stream": False,
                "options": {
                    "temperature": 0.3,
                    "num_predict": 1500
                }
            },
            timeout=180
        )
        response.raise_for_status()
        return response.json()["response"]
    except requests.exceptions.ConnectionError:
        logging.error("Kan geen verbinding maken met Ollama (chat)")
        raise Exception("Kan geen verbinding maken met Ollama. Is Ollama actief?")
    except requests.exceptions.Timeout:
        logging.error(f"Ollama chat timeout voor model {model}")
        raise Exception(f"Ollama timeout - het model reageert niet binnen 3 minuten.")
    except Exception as e:
        logging.error(f"Ollama chat fout met model {model}", exc_info=True)
        raise Exception(f"Ollama fout: {type(e).__name__}: {str(e)}")


def chat_with_openai(transcript: str, question: str, chat_history: list, api_key: str) -> str:
    """Chat about transcript using OpenAI API."""
    from openai import OpenAI

    client = OpenAI(api_key=api_key)
    chat_prompt_base_len = len(CHAT_SYSTEM_PROMPT) - len("{transcript}")
    effective_limit = get_effective_limit(TRANSCRIPT_LIMITS["openai_chat"], chat_prompt_base_len)
    system_prompt = CHAT_SYSTEM_PROMPT.format(transcript=transcript[:effective_limit])

    messages = [{"role": "system", "content": system_prompt}]
    for msg in chat_history[-10:]:
        messages.append({"role": msg["role"], "content": msg["content"]})
    messages.append({"role": "user", "content": question})

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=messages,
            temperature=0.3,
            max_tokens=1500
        )
        return response.choices[0].message.content
    except Exception as e:
        logging.error("OpenAI chat fout", exc_info=True)
        raise Exception(f"OpenAI fout: {type(e).__name__}: {str(e)}")


def chat_with_anthropic(transcript: str, question: str, chat_history: list, api_key: str) -> str:
    """Chat about transcript using Anthropic API."""
    import anthropic

    client = anthropic.Anthropic(api_key=api_key)
    chat_prompt_base_len = len(CHAT_SYSTEM_PROMPT) - len("{transcript}")
    effective_limit = get_effective_limit(TRANSCRIPT_LIMITS["anthropic_chat"], chat_prompt_base_len)
    system_prompt = CHAT_SYSTEM_PROMPT.format(transcript=transcript[:effective_limit])

    messages = []
    for msg in chat_history[-10:]:
        messages.append({"role": msg["role"], "content": msg["content"]})
    messages.append({"role": "user", "content": question})

    try:
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=1500,
            system=system_prompt,
            messages=messages
        )
        return response.content[0].text
    except Exception as e:
        logging.error("Anthropic chat fout", exc_info=True)
        raise Exception(f"Anthropic fout: {type(e).__name__}: {str(e)}")


def create_word_document(title: str, video_id: str, provider: str, model: str, summary: str) -> Document:
    """Create a Word document from the summary."""
    doc = Document()

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta = doc.add_paragraph()
    meta.add_run("Video: ").bold = True
    meta.add_run(f"https://youtube.com/watch?v={video_id}\n")
    meta.add_run("Datum: ").bold = True
    meta.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    meta.add_run("Model: ").bold = True
    meta.add_run(f"{provider}" + (f" ({model})" if model else ""))

    doc.add_paragraph()  # Spacing

    # Parse and add summary content
    lines = summary.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Heading 2 (##)
        if line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        # Heading 3 (###)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        # Bullet point
        elif line.startswith('- '):
            para = doc.add_paragraph(line[2:], style='List Bullet')
        # Numbered list
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            doc.add_paragraph(text, style='List Number')
        # Horizontal rule
        elif line.startswith('---'):
            doc.add_paragraph('─' * 50)
        # Regular paragraph
        else:
            # Handle bold text (**text**)
            para = doc.add_paragraph()
            parts = re.split(r'(\*\*[^*]+\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    para.add_run(part[2:-2]).bold = True
                else:
                    para.add_run(part)

    return doc


def chat_with_transcript(transcript: str, question: str, chat_history: list,
                         provider: str, api_key: Optional[str] = None, model: str = None) -> str:
    """Chat about a transcript using specified provider."""
    if provider == "ollama":
        return chat_with_ollama(transcript, question, chat_history, model or "gpt-oss:20b")
    elif provider == "openai":
        if not api_key:
            raise Exception("OpenAI API key is vereist.")
        return chat_with_openai(transcript, question, chat_history, api_key)
    elif provider == "anthropic":
        if not api_key:
            raise Exception("Anthropic API key is vereist.")
        return chat_with_anthropic(transcript, question, chat_history, api_key)
    else:
        raise Exception(f"Onbekende provider: {provider}")


def process_video(url: str, provider: str, api_key: Optional[str] = None,
                  model: Optional[str] = None, progress_callback=None) -> Tuple[Path, Path]:
    """
    Process a YouTube video: get transcript and create summary.
    Returns paths to transcript and summary files.
    """
    # Extract video ID
    if progress_callback:
        progress_callback("Video ID extraheren...")

    video_id = extract_video_id(url)
    if not video_id:
        raise Exception("Ongeldige YouTube URL. Controleer de link en probeer opnieuw.")

    # Get video title
    if progress_callback:
        progress_callback("Video titel ophalen...")
    title = get_video_title(video_id)

    # Create output directory
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = re.sub(r'[^\w\s-]', '', title).strip()[:50]
    base_filename = f"{timestamp}_{safe_title}"

    # Get transcript
    if progress_callback:
        progress_callback("Transcriptie ophalen van YouTube...")
    transcript, lang = get_transcript(video_id)

    # Save transcript
    transcript_path = OUTPUT_DIR / f"{base_filename}_transcriptie.txt"
    with open(transcript_path, 'w', encoding='utf-8') as f:
        f.write(f"Video: {title}\n")
        f.write(f"URL: https://youtube.com/watch?v={video_id}\n")
        f.write(f"Taal transcriptie: {lang}\n")
        f.write(f"Datum: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
        f.write("=" * 50 + "\n\n")
        f.write(transcript)

    # Create summary
    if progress_callback:
        progress_callback(f"Samenvatting maken met {provider}...")
    summary = summarize(transcript, provider, api_key, model)

    # Save summary as Word document
    summary_path = OUTPUT_DIR / f"{base_filename}_samenvatting.docx"
    doc = create_word_document(title, video_id, provider, model, summary)
    doc.save(summary_path)

    return transcript_path, summary_path


def load_config() -> dict:
    """Load configuration from file."""
    config_path = Path.home() / ".youtube_samenvatting_config.json"
    if config_path.exists():
        with open(config_path, 'r') as f:
            return json.load(f)
    return {}


def save_config(config: dict):
    """Save configuration to file."""
    config_path = Path.home() / ".youtube_samenvatting_config.json"
    with open(config_path, 'w') as f:
        json.dump(config, f)


if __name__ == "__main__":
    # CLI mode
    import sys

    if len(sys.argv) < 2:
        print("Gebruik: python youtube_samenvatting.py <youtube_url> [provider]")
        print("Providers: ollama, openai, anthropic")
        sys.exit(1)

    url = sys.argv[1]
    provider = sys.argv[2] if len(sys.argv) > 2 else "ollama"

    config = load_config()
    api_key = None
    if provider == "openai":
        api_key = config.get("openai_api_key") or os.environ.get("OPENAI_API_KEY")
    elif provider == "anthropic":
        api_key = config.get("anthropic_api_key") or os.environ.get("ANTHROPIC_API_KEY")

    try:
        print(f"Verwerken van: {url}")
        transcript_path, summary_path = process_video(
            url, provider, api_key,
            progress_callback=lambda msg: print(f"  > {msg}")
        )
        print(f"\nKlaar!")
        print(f"Transcriptie: {transcript_path}")
        print(f"Samenvatting: {summary_path}")
    except Exception as e:
        print(f"Fout: {e}")
        sys.exit(1)
